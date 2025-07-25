import streamlit as st
import requests
import os
import pandas as pd
from PIL import Image
from docx import Document
import fitz  # PyMuPDF
import numpy as np
import uuid
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime
from supabase import create_client, Client
from openai import OpenAI
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
import hashlib

# Load secrets from .streamlit/secrets.toml
client = OpenAI(api_key=st.secrets["openai_api_key"])
OCR_SPACE_API_KEY = st.secrets["ocr_space_api_key"]
SUPABASE_URL = st.secrets["supabase_url"]
SUPABASE_API_KEY = st.secrets["supabase_api_key"]
supabase: Client = create_client(SUPABASE_URL, SUPABASE_API_KEY)

GITHUB_USER = "siddhant2397"
GITHUB_REPO = "NALCObot"
GITHUB_BRANCH = "main"
GITHUB_FOLDER = "data"

EMBEDDING_MODEL = "text-embedding-3-small"
EMBEDDING_DIM = 1536

@st.cache_data(ttl=3600)
def list_github_files(user, repo, branch="main"):
    api_url = f"https://api.github.com/repos/{user}/{repo}/contents/{GITHUB_FOLDER}"
    response = requests.get(api_url)
    if response.status_code != 200:
        st.error(f"GitHub API error: {response.status_code} - {response.text}")
        return []
    try:
        files = response.json()
        st.write("🔎 Raw response from GitHub API:", files)
        # Ensure it's a list of dicts with 'name' key
        valid_files = [
            f for f in files
            if isinstance(f, dict)
            and "name" in f
            and f["name"].endswith((".docx", ".xlsx", ".pdf"))
        ]
        return valid_files
    except Exception as e:
        st.error(f"Failed to parse GitHub file list: {e}")
        return []

        

def fetch_file(file):
    raw_url = f"https://raw.githubusercontent.com/{GITHUB_USER}/{GITHUB_REPO}/{GITHUB_BRANCH}/{GITHUB_FOLDER}/{file['name']}"
    response = requests.get(raw_url)
    with open(file['name'], "wb") as f:
        f.write(response.content)
    return file['name']

def ocr_space_image(filepath):
    with open(filepath, 'rb') as f:
        response = requests.post(
            url='https://api.ocr.space/parse/image',
            files={'file': f},
            data={'apikey': OCR_SPACE_API_KEY}
        )
    result = response.json()
    return result['ParsedResults'][0]['ParsedText'] if 'ParsedResults' in result else ""

def extract_text(filepath):
    if filepath.endswith(".docx"):
        return extract_docx_text(filepath)
    elif filepath.endswith(".xlsx"):
        df = pd.read_excel(filepath)
        return df.to_string(index=False)
    elif filepath.endswith(".pdf"):
        return extract_pdf_text(filepath)
    return ""

def extract_docx_text(filepath):
    doc = Document(filepath)
    text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            text += "\n" + row_text
    for rel in doc.part._rels:
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            img_data = rel.target_part.blob
            with open("temp.png", "wb") as f:
                f.write(img_data)
            text += "\n" + ocr_space_image("temp.png")
            os.remove("temp.png")
    return text

def extract_pdf_text(filepath):
    doc = fitz.open(filepath)
    text = " ".join([page.get_text() for page in doc])
    if len(text.strip()) < 100:
        st.info(f"🔍 Attempting OCR on scanned PDF: {filepath}")
        images = [page.get_pixmap().pil_tobytes(format="PNG") for page in doc]
        for i, img_bytes in enumerate(images):
            with open(f"page_{i}.png", "wb") as f:
                f.write(img_bytes)
        text = ""
        for i in range(len(images)):
            text += ocr_space_image(f"page_{i}.png") + "\n"
            os.remove(f"page_{i}.png")
    return text

def embed_text(text):
    response = client.embeddings.create(
        input=text,
        model=EMBEDDING_MODEL
    )
    return np.array(response.data[0].embedding, dtype=np.float32)



def store_embeddings(chunks):
    count = 0

    # Step 1: Fetch existing hashes from Supabase
    existing = supabase.table("documents").select("hash").execute()
    existing_hashes = {row["hash"] for row in existing.data}

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue

        # Step 2: Compute hash of the chunk
        chunk_hash = hashlib.sha256(chunk.encode("utf-8")).hexdigest()

        # Step 3: Skip if already present
        if chunk_hash in existing_hashes:
            continue

        # Step 4: Embed and store if new
        emb = embed_text(chunk).tolist()
        supabase.table("documents").insert({
            "id": str(uuid.uuid4()),
            "hash": chunk_hash,
            "content": chunk,
            "embedding": json.dumps(emb)
        }).execute()
        count += 1

    st.info(f"📦 Stored {count} new chunks in Supabase.")


def log_interaction(question, input_tokens, output_tokens, cost):
    supabase.table("chat_logs").insert({
        "id": str(uuid.uuid4()),
        "timestamp": datetime.utcnow().isoformat(),
        "question": question,
        "prompt_tokens": input_tokens,
        "completion_tokens": output_tokens,
        "cost": round(cost, 6)
    }).execute()

def get_monthly_usage():
    response = supabase.table("chat_logs").select("*").execute()
    rows = response.data
    current_month = datetime.utcnow().month
    prompt_sum = sum(r['prompt_tokens'] for r in rows if datetime.fromisoformat(r['timestamp']).month == current_month)
    completion_sum = sum(r['completion_tokens'] for r in rows if datetime.fromisoformat(r['timestamp']).month == current_month)
    cost_sum = sum(r['cost'] for r in rows if datetime.fromisoformat(r['timestamp']).month == current_month)
    return prompt_sum, completion_sum, cost_sum

def get_top_chunks(question_embedding, k=3):
    response = supabase.table("documents").select("content", "embedding").execute()
    results = response.data

    if not results:
        st.warning("⚠️ No document embeddings found in the database.")
        return []

    chunks = [r['content'] for r in results]
    embeddings = []
    for r in results:
        if r['embedding']:
            try:
                emb_list = r['embedding']
                if isinstance(emb_list, str):
                    emb_list = json.loads(emb_list)
                embeddings.append(np.array(emb_list, dtype=np.float32))
            except Exception as e:
                st.warning(f"⚠️ Skipping corrupted embedding: {e}")

    



    if not embeddings:
        st.warning("⚠️ No valid embeddings available.")
        return []

    similarities = cosine_similarity([question_embedding], embeddings)[0]
    top_k_idx = similarities.argsort()[-k:][::-1]
    return [chunks[i] for i in top_k_idx]


def split_text(text, chunk_size=500, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    return splitter.split_text(text)


def ask_gpt(question, context):
    prompt = f"Answer the following question based on the context:\n\nContext:\n{context}\n\nQuestion: {question}"
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=400
    )
    answer = response.choices[0].message.content
    usage = response.usage
    input_tokens = usage.prompt_tokens
    output_tokens = usage.completion_tokens
    cost = (input_tokens / 1000 * 0.0005) + (output_tokens / 1000 * 0.0015)
    return answer, input_tokens, output_tokens, cost

st.title("📄 CISF Chat Bot")

question = st.text_input("Ask your question")

if question:
    st.info("📂 Processing documents")
    files = list_github_files(GITHUB_USER, GITHUB_REPO, GITHUB_BRANCH)
    all_chunks = []

    for file in files:
        local_name = fetch_file(file)
        st.write(f"📄 Processing: {file['name']}")
        text = extract_text(local_name)
        st.write(f"📏 Extracted length: {len(text)} characters")
        chunks = split_text(text)
        st.write(f"🧩 Generated {len(chunks)} chunks")
        all_chunks.extend(chunks)
        os.remove(local_name)

    if all_chunks:
        st.success(f"✅ Loaded {len(files)} files with {len(all_chunks)} chunks.")
        store_embeddings(all_chunks)

        question_embedding = embed_text(question)
        top_chunks = get_top_chunks(question_embedding, k=3)

        if not top_chunks:
            st.stop()

        context = "\n---\n".join(top_chunks)
        st.info("💬 Generating answer")
        answer, input_tokens, output_tokens, cost = ask_gpt(question, context)
        log_interaction(question, input_tokens, output_tokens, cost)

        st.success("✅ Answer:")
        st.write(answer)

        st.markdown("---")
        st.subheader("📊 Cost Summary for This Query")
        st.write(f"**Prompt tokens:** {input_tokens}")
        st.write(f"**Response tokens:** {output_tokens}")
        st.write(f"**Estimated cost:** ${cost:.6f} USD")

        prompt_sum, completion_sum, cost_sum = get_monthly_usage()
        st.markdown("---")
        st.subheader("📆 Monthly Usage Summary")
        st.write(f"**Total input tokens:** {prompt_sum}")
        st.write(f"**Total output tokens:** {completion_sum}")
        st.write(f"**Estimated monthly cost:** ${cost_sum:.4f} USD")
    else:
        st.warning("⚠️ No usable content was extracted from documents. Please check if the files contain readable text.")
